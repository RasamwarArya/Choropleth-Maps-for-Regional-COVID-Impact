"""
Script to convert markdown documentation to Word format (.docx)
"""

import pypandoc
import os

def convert_markdown_to_word():
    """Convert the COVID documentation from markdown to Word format"""
    
    # Input and output file paths
    input_file = "COVID_Data_Visualization_Documentation.md"
    output_file = "COVID_Data_Visualization_Documentation.docx"
    
    try:
        print("Converting markdown to Word format...")
        
        # Convert markdown to docx
        pypandoc.convert_file(
            input_file, 
            'docx', 
            outputfile=output_file,
            extra_args=[
                '--reference-doc=reference.docx' if os.path.exists('reference.docx') else None,
                '--toc',  # Add table of contents
                '--toc-depth=3',  # TOC depth
                '--number-sections',  # Number sections
                '--highlight-style=github'  # Code highlighting
            ]
        )
        
        print(f"✅ Successfully converted to: {output_file}")
        print(f"📄 File size: {os.path.getsize(output_file)} bytes")
        
        return True
        
    except Exception as e:
        print(f"❌ Error during conversion: {e}")
        
        # Fallback method - create a simple Word document
        print("🔄 Trying alternative method...")
        return create_simple_word_doc()

def create_simple_word_doc():
    """Create a simple Word document using python-docx"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.style import WD_STYLE_TYPE
        
        print("Creating Word document using python-docx...")
        
        # Read the markdown content
        with open("COVID_Data_Visualization_Documentation.md", 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('COVID-19 Data Visualization Project: Comprehensive Documentation', 0)
        
        # Split content into sections and convert
        lines = content.split('\n')
        current_section = ""
        
        for line in lines:
            line = line.strip()
            
            if line.startswith('# '):
                # Main heading
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # Sub heading
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # Sub-sub heading
                doc.add_heading(line[4:], level=3)
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                p.add_run(line[2:-2]).bold = True
            elif line.startswith('• ') or line.startswith('- '):
                # Bullet point
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('✅') or line.startswith('❌'):
                # Special bullet points
                doc.add_paragraph(line, style='List Bullet')
            elif line.startswith('```'):
                # Code block (skip for now)
                continue
            elif line and not line.startswith('---'):
                # Regular paragraph
                if line:
                    doc.add_paragraph(line)
        
        # Save the document
        output_file = "COVID_Data_Visualization_Documentation.docx"
        doc.save(output_file)
        
        print(f"✅ Successfully created: {output_file}")
        print(f"📄 File size: {os.path.getsize(output_file)} bytes")
        
        return True
        
    except ImportError:
        print("❌ python-docx not installed. Installing...")
        os.system("pip install python-docx")
        return create_simple_word_doc()
    except Exception as e:
        print(f"❌ Error creating Word document: {e}")
        return False

def create_conversion_instructions():
    """Create instructions for manual conversion"""
    instructions = """
# Manual Conversion Instructions

## Method 1: Using Microsoft Word
1. Open Microsoft Word
2. Go to File > Open
3. Select "COVID_Data_Visualization_Documentation.md"
4. Word will automatically convert the markdown
5. Save as .docx format

## Method 2: Using Online Converters
1. Go to https://pandoc.org/try/ or https://www.markdowntoword.com/
2. Upload your markdown file
3. Download the converted Word document

## Method 3: Copy and Paste
1. Open the markdown file in a text editor
2. Copy all content
3. Paste into Microsoft Word
4. Apply formatting manually:
   - Make headings bold and larger
   - Add bullet points for lists
   - Format code blocks with monospace font

## Method 4: Using Google Docs
1. Upload markdown file to Google Drive
2. Open with Google Docs
3. Download as .docx format
"""
    
    with open("Word_Conversion_Instructions.txt", 'w') as f:
        f.write(instructions)
    
    print("📝 Created Word_Conversion_Instructions.txt")

if __name__ == "__main__":
    print("🔄 Converting COVID Documentation to Word Format")
    print("=" * 50)
    
    success = convert_markdown_to_word()
    
    if not success:
        print("\n📋 Manual conversion methods available:")
        create_conversion_instructions()
        print("Check Word_Conversion_Instructions.txt for step-by-step guide")
